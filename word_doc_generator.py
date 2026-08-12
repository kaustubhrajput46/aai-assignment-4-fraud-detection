import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def create_apa_doc(python_code, image_path, output_path):
    doc = Document()
    
    # Font settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Margin settings
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Header setup (except first page usually, but here we add it for all to keep it simple, or we can handle it)
        # Actually APA says page numbers on all pages including title page top right
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run()
        add_page_number(run)

    # Title Page
    for _ in range(3):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Use Unsupervised Deep Learning Algorithm to Detect Fraud with PyOD")
    title_run.bold = True
    
    doc.add_paragraph() # extra space
    
    details = [
        "Kaustubh Rajput",
        "University of the Cumberlands",
        "MSCS-633-B01 Advance Artificial Intelligence",
        "Dr. V",
        "July 21, 2026"
    ]
    
    for detail in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(detail)
        
    doc.add_page_break()
    
    # Level 1 Heading
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = h1.add_run("Use Unsupervised Deep Learning Algorithm to Detect Fraud with PyOD")
    r1.bold = True
    
    p = doc.add_paragraph("This assignment implements an unsupervised deep learning approach using the PyOD library's AutoEncoder to detect credit card fraud. Due to dataset constraints, a synthetic dataset mimicking the structure of the standard Kaggle creditcard.csv dataset was generated.")
    p.paragraph_format.line_spacing = 2.0
    
    p2 = doc.add_paragraph("The code can be found at the following GitHub repository: ")
    p2.add_run("https://github.com/kaustubhrajput46/aai-fraud-detection")
    p2.paragraph_format.line_spacing = 2.0
    
    # Level 2 Heading
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = h2.add_run("Python Source Code")
    r2.bold = True
    
    # Code block
    code_p = doc.add_paragraph()
    code_r = code_p.add_run(python_code)
    code_r.font.name = 'Courier New'
    code_r.font.size = Pt(10)
    
    # Level 2 Heading
    h3 = doc.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r3 = h3.add_run("Output and Visualization")
    r3.bold = True
    
    p3 = doc.add_paragraph("The AutoEncoder was trained on the synthetic normal and fraudulent transactions. The reconstruction error scores were calculated and are displayed in the histogram below. The classification report printed in the terminal shows the precision, recall, and F1 scores of the anomaly detection model.")
    p3.paragraph_format.line_spacing = 2.0
    
    doc.add_picture(image_path, width=Inches(6.0))
    fig_p = doc.add_paragraph("Figure 1. Distribution of Outlier Scores (Reconstruction Error) for Normal and Fraudulent Transactions.")
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p.runs[0].italic = True

    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    with open('fraud_detection.py', 'r') as f:
        code = f.read()
    create_apa_doc(code, 'output.png', 'Kaustubh_Rajput_AAI_Assignment_4.docx')
